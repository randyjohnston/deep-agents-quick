"""Typed DOCX generation and bounded text extraction."""

from __future__ import annotations

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from pydantic import BaseModel, Field

from app.tools.office.paths import (
    load_office_file,
    load_office_template,
    resolve_read_path,
    resolve_template_path,
    resolve_write_path,
)
from app.tools.office.theme import ResolvedTheme, Theme, resolve_theme

_DOTX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
)
_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
)


class Section(BaseModel):
    """A headed document section containing paragraphs and optional bullets."""

    heading: str | None = Field(default=None, description="Section heading")
    paragraphs: list[str] = Field(default_factory=list, description="Body paragraphs")
    bullets: list[str] = Field(default_factory=list, description="Bulleted list items")


class Document(BaseModel):
    """The structured content of one Word document."""

    title: str | None = Field(default=None, description="Document title")
    subtitle: str | None = Field(default=None, description="Optional subtitle")
    sections: list[Section] = Field(default_factory=list)


def write_docx(
    filename: str,
    document: Document,
    theme: Theme | None = None,
    template: str | None = None,
) -> str:
    """Create a Word .docx file from typed content and return its path.

    Args:
        filename: Output basename. The .docx extension is added when absent.
        document: Title, optional subtitle, and headed body sections.
        theme: Bounded inline branding or a named theme reference.
        template: Optional .dotx/.docx template under OFFICE_INPUT_DIR.
    """
    has_body = any(s.heading or s.paragraphs or s.bullets for s in document.sections)
    if not document.title and not document.subtitle and not has_body:
        raise ValueError("Provide a title, subtitle, or at least one section.")
    path = resolve_write_path(filename, ".docx")
    path.parent.mkdir(parents=True, exist_ok=True)
    template_path = resolve_template_path(template, (".dotx", ".docx")) if template else None
    output = (
        load_office_template(template_path, DocxDocument, _DOTX_CONTENT_TYPE, _DOCX_CONTENT_TYPE)
        if template_path
        else DocxDocument()
    )
    branding = resolve_theme(theme)
    _apply_theme(output, branding)
    if branding and branding.logo:
        output.add_picture(str(branding.logo), width=Inches(1.5))
    if document.title:
        output.add_heading(document.title, level=0)
    if document.subtitle:
        output.add_paragraph(document.subtitle, style="Subtitle")
    for section in document.sections:
        if section.heading:
            output.add_heading(section.heading, level=1)
        for paragraph in section.paragraphs:
            output.add_paragraph(paragraph)
        for bullet in section.bullets:
            output.add_paragraph(bullet, style="List Bullet")
    output.save(path)
    return f"Wrote {path} — {len(document.sections)} sections"


def _apply_theme(output, theme: ResolvedTheme | None) -> None:
    if not theme:
        return
    normal = output.styles["Normal"]
    if theme.body_font:
        normal.font.name = theme.body_font
    for name in ("Title", "Heading 1"):
        style = output.styles[name]
        if theme.heading_font:
            style.font.name = theme.heading_font
        if theme.header_background:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), theme.header_background)
            style.element.get_or_add_pPr().append(shading)
        color = theme.header_foreground or theme.accent_color
        if color:
            style.font.color.rgb = RGBColor.from_string(color)


def read_docx(path: str, max_blocks: int = 500) -> str:
    """Extract paragraph and table text from a permitted .docx file."""
    if max_blocks < 1:
        raise ValueError("max_blocks must be at least 1.")
    resolved = resolve_read_path(path, (".docx",))
    document = load_office_file(resolved, DocxDocument)
    values = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        values.extend(
            " | ".join(cell.text for cell in row.cells)
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        )
    blocks = [f"# {resolved}", "", *(values[:max_blocks] or ["(empty)"])]
    if len(values) > max_blocks:
        blocks.append(f"... truncated at {max_blocks} blocks; raise max_blocks for more")
    return "\n".join(blocks)
