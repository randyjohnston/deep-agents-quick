from __future__ import annotations

import pytest
from docx import Document as DocxDocument

from app.tools.office.docx import Document, Section, read_docx, write_docx


def test_docx_round_trip_with_structure(_isolate_dirs):
    result = write_docx(
        "brief",
        Document(
            title="Market Brief",
            subtitle="August 2026",
            sections=[Section(heading="Findings", paragraphs=["Demand grew."], bullets=["North"])]
        ),
    )
    path = _isolate_dirs / "out" / "brief.docx"
    assert path.is_file(), result
    parsed = DocxDocument(path)
    assert [p.text for p in parsed.paragraphs] == [
        "Market Brief", "August 2026", "Findings", "Demand grew.", "North"
    ]
    assert "Demand grew." in read_docx("brief.docx")


def test_docx_rejects_empty_content(_isolate_dirs):
    with pytest.raises(ValueError, match="Provide"):
        write_docx("empty", Document(sections=[Section()]))


def test_docx_rejects_macro_extension(_isolate_dirs):
    content = Document(title="Safe")
    with pytest.raises(ValueError, match="Macro-enabled"):
        write_docx("unsafe.docm", content)
